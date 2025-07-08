import os, re, requests, fitz, pytesseract, cv2, numpy as np
from io import BytesIO
from bs4 import BeautifulSoup
from urllib.parse import urlparse, unquote
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from PIL import Image
from dotenv import load_dotenv
from utils.link_check import extract_links_from_any_url

load_dotenv()

# Optional: set path if Tesseract is not in PATH (for Windows)
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
def clean_text_for_summary(raw_text):
    text = raw_text

    # Remove headers like "[Page 1 OCR]"
    text = re.sub(r'\[Page \d+ (OCR|Text)\]', '', text)

    # Remove non-ASCII characters (from OCR garbage)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)

    # Remove repeated punctuation, like -------- or =====
    text = re.sub(r'[-=]{4,}', ' ', text)

    # Remove duplicate lines
    lines = text.splitlines()
    seen = set()
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        if line and line not in seen:
            seen.add(line)
            cleaned_lines.append(line)

    cleaned = "\n".join(cleaned_lines)

    # Remove multiple blank lines and collapse whitespace
    cleaned = re.sub(r'\n{2,}', '\n', cleaned)
    cleaned = re.sub(r'[ \t]{2,}', ' ', cleaned)

    return cleaned.strip()

def extract_text(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}

        # Google Drive link support
        if "drive.google.com" in url:
            match = re.search(r"/d/([a-zA-Z0-9_-]+)", url) or re.search(r"id=([a-zA-Z0-9_-]+)", url)
            if match:
                file_id = match.group(1)
                url = f"https://drive.google.com/uc?export=download&id={file_id}"

        # Download file
        r = requests.get(url, headers=headers, timeout=30)
        r.raise_for_status()
        content = BytesIO(r.content)
        content_type = r.headers.get("Content-Type", "").lower()
        url_lower = url.lower()

        # ---------- PDF ----------
        if url_lower.endswith(".pdf") or "pdf" in content_type:
            text = ""
            pdf = fitz.open(stream=content, filetype="pdf")
            for page in pdf:
                t = page.get_text("text").strip()
                if t and len(t) > 50:
                    text += "\n" + t
                else:
                    pix = page.get_pixmap(matrix=fitz.Matrix(3, 3))
                    img = Image.open(BytesIO(pix.tobytes("png"))).convert("RGB")
                    img_array = np.array(img)
                    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
                    ocr_text = pytesseract.image_to_string(gray, config="--psm 6").strip()
                    text += "\n" + ocr_text
            return text.strip()

        # ---------- DOCX ----------
        elif url_lower.endswith(".docx") or "word" in content_type:
            import docx
            doc = docx.Document(content)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()

        # ---------- TXT ----------
        elif url_lower.endswith(".txt") or "text/plain" in content_type:
            return content.read().decode('utf-8', errors='ignore').strip()

        # ---------- IMAGE (PNG, JPG) ----------
        elif url_lower.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff")) or "image" in content_type:
            img = Image.open(content).convert("RGB")
            img_np = np.array(img)
            gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
            
            # Enhance image for better OCR
            # Apply some preprocessing to improve OCR accuracy
            gray = cv2.GaussianBlur(gray, (1, 1), 0)
            gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            
            # Try multiple OCR configurations for better results
            configs = [
                "--psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz ",
                "--psm 4",
                "--psm 3", 
                "--psm 6",
                "--psm 1",
                "--psm 8",
                "--psm 13"
            ]
            
            best_text = ""
            for config in configs:
                try:
                    text = pytesseract.image_to_string(gray, config=config).strip()
                    if len(text) > len(best_text):
                        best_text = text
                        if len(best_text) > 50:  # If we get good results, break early
                            break
                except Exception as e:
                    continue
            
            # If OCR fails or returns very little text, create a descriptive fallback
            if len(best_text) < 5:
                best_text = f"Image from {url}. This appears to be an image file. OCR was unable to extract readable text content. The image may contain graphics, charts, or text that is not clearly readable by OCR technology."
            
            print(f"📸 OCR extracted {len(best_text)} characters from image")
            return best_text

        # ---------- WEB PAGE ----------
        else:
            r.encoding = 'utf-8'
            soup = BeautifulSoup(r.text, 'html.parser')
            for tag in soup(['script', 'style', 'nav', 'footer', 'noscript']):
                tag.decompose()
            return re.sub(r'\n{2,}', '\n', soup.get_text(separator='\n', strip=True)).strip()

    except Exception as e:
        print(f"❌ {url} failed: {e}")
        # Return a basic description instead of empty string
        return f"Failed to extract content from {url}: {str(e)}"


def build_vectordb(links):
    print(f"🔍 Building vector DB for {len(links)} links: {links}")
    
    try:
        if "drive.google.com" in links[0]:
            match = re.search(r"/d/([a-zA-Z0-9_-]+)", links[0]) or re.search(r"id=([a-zA-Z0-9_-]+)", links[0])
            name = match.group(1) if match else "drivefile"
        else:
            path = unquote(urlparse(links[0]).path)
            filename = os.path.basename(path)
            name = filename.split('.')[0] if '.' in filename else filename or "file"
    except Exception:
        name = "file"

    vectordb_path = f"vectordb/{name}_vectordb"
    os.makedirs(vectordb_path, exist_ok=True)

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    docs, full_text = [], ""

    for i, link in enumerate(links[:50]):
        print(f"🔗 Processing {i+1}/{len(links)}: {link}")
        text = extract_text(link)
        print(f"📝 Extracted {len(text)} characters")
        if len(text) > 20:  # Show preview for substantial content
            print(f"📋 Preview: {text[:100]}...")
        
        if text:
            full_text += text + "\n\n"
            chunks = splitter.split_text(text)
            if not chunks:
                chunks = [text]  # fallback for short content
            for chunk in chunks:
                docs.append(Document(page_content=chunk, metadata={"source": link}))

    # Always create at least one document, even if content is minimal
    if not docs:
        fallback_text = f"Content processed from {len(links)} link(s). No readable text could be extracted."
        docs.append(Document(page_content=fallback_text, metadata={"source": links[0] if links else "unknown"}))
        full_text = fallback_text
        print("⚠️ No valid content found. Created fallback document.")

    emb = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    Chroma.from_documents(docs, emb, persist_directory=vectordb_path)
    print(f"✅ Vector DB created at: {vectordb_path} with {len(docs)} chunks")
    return vectordb_path, full_text